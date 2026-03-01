"""
Contract Export Service
Export draft kontrak ke Word (.docx) dan PDF
"""
import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY


def export_contract_docx(title: str, content: str) -> bytes:
    """Export contract draft to Word (.docx) format."""
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # -- Parse content and build document --
    lines = content.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Detect headings: lines that are ALL CAPS or start with "PASAL" or "BAB"
        is_heading = (
            stripped.isupper() and len(stripped) > 3
            or stripped.upper().startswith("PASAL ")
            or stripped.upper().startswith("BAB ")
        )

        # Detect title-like lines (centered, bold)
        is_title = (
            stripped.upper().startswith("PERJANJIAN")
            or stripped.upper().startswith("SURAT")
            or stripped.upper().startswith("KONTRAK")
            or stripped.upper().startswith("NON-DISCLOSURE")
            or stripped.upper().startswith("MEMORANDUM")
            or (stripped.startswith("Nomor") or stripped.startswith("NOMOR"))
        )

        if is_title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(26, 26, 46)
        elif is_heading:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(22, 33, 62)
            p.space_before = Pt(12)
            p.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            # Handle bold markers **text**
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
            p.paragraph_format.line_spacing = Pt(16)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_contract_pdf(title: str, content: str) -> bytes:
    """Export contract draft to PDF format."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=25 * mm,
        leftMargin=25 * mm,
        topMargin=25 * mm,
        bottomMargin=25 * mm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "ContractTitle",
        parent=styles["Title"],
        fontSize=16,
        textColor=colors.HexColor("#1a1a2e"),
        alignment=TA_CENTER,
        spaceAfter=6,
        fontName="Helvetica-Bold",
    )

    heading_style = ParagraphStyle(
        "ContractHeading",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#16213e"),
        alignment=TA_CENTER,
        spaceBefore=14,
        spaceAfter=6,
        fontName="Helvetica-Bold",
    )

    body_style = ParagraphStyle(
        "ContractBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=15,
        alignment=TA_JUSTIFY,
        fontName="Helvetica",
    )

    story = []
    lines = content.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 3 * mm))
            continue

        is_heading = (
            stripped.isupper() and len(stripped) > 3
            or stripped.upper().startswith("PASAL ")
            or stripped.upper().startswith("BAB ")
        )

        is_title = (
            stripped.upper().startswith("PERJANJIAN")
            or stripped.upper().startswith("SURAT")
            or stripped.upper().startswith("KONTRAK")
            or stripped.upper().startswith("NON-DISCLOSURE")
            or stripped.upper().startswith("MEMORANDUM")
            or stripped.startswith("Nomor") or stripped.startswith("NOMOR")
        )

        # Sanitize for ReportLab (escape HTML-like chars)
        safe_text = (
            stripped
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        # Convert **bold** markers
        safe_text = re.sub(
            r'\*\*(.*?)\*\*',
            r'<b>\1</b>',
            safe_text,
        )

        if is_title:
            story.append(Paragraph(safe_text, title_style))
        elif is_heading:
            story.append(Paragraph(safe_text, heading_style))
        else:
            story.append(Paragraph(safe_text, body_style))

    # Footer
    story.append(Spacer(1, 15 * mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
    footer_style = ParagraphStyle(
        "Footer",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.grey,
        alignment=TA_CENTER,
    )
    story.append(Spacer(1, 3 * mm))
    story.append(
        Paragraph(
            "Dokumen ini dihasilkan oleh LexAI — AI Copilot untuk Praktisi Hukum Indonesia. "
            "Draft kontrak ini bersifat referensi dan disarankan untuk dikonsultasikan dengan ahli hukum.",
            footer_style,
        )
    )

    doc.build(story)
    return buffer.getvalue()
