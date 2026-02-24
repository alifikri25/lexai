"""
Document Parser Service
Extracts text from PDF and DOCX files using PyMuPDF and python-docx
"""
import base64
import io
import fitz  # PyMuPDF
from docx import Document


def parse_pdf(file_bytes: bytes) -> dict:
    """Parse PDF file and extract text content."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    total_pages = len(doc)
    
    for page_num in range(total_pages):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            text_parts.append(f"--- Halaman {page_num + 1} ---\n{text}")
    
    doc.close()
    
    full_text = "\n\n".join(text_parts)
    
    # Limit text to prevent token overflow (max ~50000 chars for LLM)
    if len(full_text) > 50000:
        full_text = full_text[:50000] + "\n\n[... Dokumen terpotong karena terlalu panjang ...]"
    
    return {
        "text": full_text,
        "pages": total_pages,
    }


def parse_docx(file_bytes: bytes) -> dict:
    """Parse DOCX file and extract text content."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    
    full_text = "\n\n".join(text_parts)
    
    # Limit text to prevent token overflow
    if len(full_text) > 50000:
        full_text = full_text[:50000] + "\n\n[... Dokumen terpotong karena terlalu panjang ...]"
    
    return {
        "text": full_text,
        "pages": 1,
    }


def parse_document(file_base64: str, file_type: str) -> dict:
    """Parse document from base64 encoded content."""
    file_bytes = base64.b64decode(file_base64)
    
    if file_type.lower() == "pdf":
        return parse_pdf(file_bytes)
    elif file_type.lower() in ["docx", "doc"]:
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Tipe file tidak didukung: {file_type}")
